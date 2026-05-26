import argparse
import json
import re
from pathlib import Path
from docx import Document

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': W_NS}
TAG_RE = re.compile(r'^\[([^\]]+)\]\s?(.*)$')


def load_translations(path):
    data = {}
    with open(path, encoding='utf-8') as handle:
        for line in handle:
            match = TAG_RE.match(line.rstrip('\n'))
            if match:
                data[match.group(1)] = match.group(2)
    return data


def set_paragraph_text(paragraph, text):
    """Replace text content preserving images and other non-text elements.

    Operates at the XML level on <w:t> elements only, so <w:drawing>,
    <w:pict>, <mc:AlternateContent>, and other non-text content within
    runs is never touched.
    """
    t_elements = paragraph._element.findall(f'.//{{{W_NS}}}t')
    if not t_elements:
        # No text elements — look for runs to insert text into
        runs = paragraph._element.findall(f'{{{W_NS}}}r')
        if runs:
            from docx.oxml import OxmlElement
            new_t = OxmlElement('w:t')
            new_t.text = text
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            runs[0].append(new_t)
        else:
            # No runs either — fallback (destructive but nothing to preserve)
            paragraph.text = text
        return

    for i, t_elem in enumerate(t_elements):
        if i == 0:
            t_elem.text = text
        else:
            t_elem.text = ''


def set_xml_paragraph_text(p_elem, text):
    runs = p_elem.findall(f'{{{W_NS}}}r')
    if not runs:
        return
    first_t = runs[0].find(f'{{{W_NS}}}t')
    if first_t is None:
        first_t = runs[0].makeelement(f'{{{W_NS}}}t')
        runs[0].append(first_t)
    first_t.text = text
    first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    for run in runs[1:]:
        for t_elem in run.findall(f'{{{W_NS}}}t'):
            t_elem.text = ''


def _find_all_txbx_elements(parent):
    """Recursively collect all w:txbxContent elements.

    MUST stay in sync with extract_docx.py:_find_all_txbx_elements to
    preserve textbox_index alignment between extraction and rendering.
    Uses explicit element iteration (not XPath .//) for robustness across
    lxml versions and namespace edge cases.
    """
    txbx_list = []
    if not hasattr(parent, 'tag') or callable(parent.tag):
        return txbx_list
    tag_str = parent.tag if isinstance(parent.tag, str) else ''
    if tag_str.endswith('}txbxContent'):
        txbx_list.append(parent)
    for child in parent:
        txbx_list.extend(_find_all_txbx_elements(child))
    return txbx_list


def build_textbox_index(doc):
    xml_bodies = [doc.element.body]
    for section in doc.sections:
        for hf in [section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer]:
            if hf is not None:
                xml_bodies.append(hf._element)
    textboxes = []
    for xml_body in xml_bodies:
        textboxes.extend(_find_all_txbx_elements(xml_body))
    return textboxes


def main():
    parser = argparse.ArgumentParser(description='Render tagged translation into DOCX template.')
    parser.add_argument('--template', required=True)
    parser.add_argument('--extraction', required=True)
    parser.add_argument('--translation', required=True)
    parser.add_argument('--output-docx', required=True)
    parser.add_argument('--render-log', required=True)
    args = parser.parse_args()

    translations = load_translations(args.translation)
    extraction = json.loads(Path(args.extraction).read_text(encoding='utf-8'))
    doc = Document(args.template)
    textboxes = build_textbox_index(doc)
    applied = []
    missing = []
    layout_warnings = []

    for entry in extraction['paragraphs']:
        tag = entry['tag']
        if tag not in translations:
            missing.append(tag)
            continue
        text = translations[tag]
        typ = entry['type']
        if typ == 'body':
            set_paragraph_text(doc.paragraphs[entry['para_index']], text)
        elif typ == 'body_table':
            cell = doc.tables[entry['table_idx']].rows[entry['row']].cells[entry['col']]
            if entry.get('para_index', 0) >= len(cell.paragraphs):
                layout_warnings.append(f"{tag}: paragraph index exceeded cell paragraph count")
                set_paragraph_text(cell.paragraphs[0], text)
            else:
                set_paragraph_text(cell.paragraphs[entry.get('para_index', 0)], text)
        elif typ == 'header':
            section = doc.sections[entry['section_idx']]
            hf = getattr(section, entry['header_type'])
            set_paragraph_text(hf.paragraphs[entry['para_index']], text)
        elif typ == 'footer':
            section = doc.sections[entry['section_idx']]
            hf = getattr(section, entry['footer_type'])
            set_paragraph_text(hf.paragraphs[entry['para_index']], text)
        elif typ == 'header_table':
            section = doc.sections[entry['section_idx']]
            hf = getattr(section, entry['header_type'])
            cell = hf.tables[entry['table_idx']].rows[entry['row']].cells[entry['col']]
            set_paragraph_text(cell.paragraphs[min(entry.get('para_index', 0), len(cell.paragraphs) - 1)], text)
        elif typ == 'footer_table':
            section = doc.sections[entry['section_idx']]
            hf = getattr(section, entry['footer_type'])
            cell = hf.tables[entry['table_idx']].rows[entry['row']].cells[entry['col']]
            set_paragraph_text(cell.paragraphs[min(entry.get('para_index', 0), len(cell.paragraphs) - 1)], text)
        elif typ == 'textbox':
            txbx = textboxes[entry['txbx_index']]
            p_elem = txbx.findall(f'{{{W_NS}}}p')[entry['para_index']]
            set_xml_paragraph_text(p_elem, text)
        applied.append(tag)

    Path(args.output_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output_docx)
    render_log = {
        'logical_tags_applied': len(applied),
        'missing': missing,
        'output_path': args.output_docx,
        'layout_warnings': layout_warnings,
    }
    Path(args.render_log).write_text(json.dumps(render_log, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"Applied {len(applied)} logical tags; missing {len(missing)}")


if __name__ == '__main__':
    main()
