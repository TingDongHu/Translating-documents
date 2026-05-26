"""Generate a DOCX test file with various text box nesting patterns."""
from docx import Document
from docx.shared import Pt, Inches
from lxml import etree
import os

DOCX_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = f'{{{DOCX_NS}}}'

def make_element(tag, text=None):
    tagname = f'{W}{tag}'
    elem = etree.SubElement(etree.Element('_dummy'), tagname)
    parent = elem.getparent()
    parent.remove(elem)
    if text is not None:
        t_elem = etree.SubElement(elem, f'{W}t')
        t_elem.text = text
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return elem

def make_run(text):
    r = etree.Element(f'{W}r')
    t = etree.SubElement(r, f'{W}t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r

def make_paragraph(text):
    p = etree.Element(f'{W}p')
    r = make_run(text)
    p.append(r)
    return p

def make_textbox_paragraph(text):
    """Create a paragraph inside w:txbxContent."""
    txbx = etree.Element(f'{W}txbxContent')
    p = make_paragraph(text)
    txbx.append(p)
    return txbx

def main():
    doc = Document()

    # --- Section 1: Normal text ---
    doc.add_paragraph("Normal paragraph — not a text box")

    # --- Section 2: Direct text box (common pattern) ---
    p2 = doc.add_paragraph()
    run2 = p2.add_run()
    r_elem2 = run2._element
    txbx2 = make_textbox_paragraph("TextBox 1: Direct text box in body")
    r_elem2.append(txbx2)

    # --- Section 3: Text box via mc:AlternateContent > wps:wsp > wps:txbx ---
    # This simulates the Word 2010+ group shape text box pattern
    MC = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    A = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    p3 = doc.add_paragraph()
    run3 = p3.add_run()
    r_elem3 = run3._element

    alt = etree.Element(f'{{{MC}}}AlternateContent')
    choice = etree.SubElement(alt, f'{{{MC}}}Choice', Requires='wps')
    wsp = etree.SubElement(choice, f'{{{WPS}}}wsp')
    wps_txbx = etree.SubElement(wsp, f'{{{WPS}}}txbx')
    w_txbx = etree.SubElement(wps_txbx, f'{W}txbxContent')
    p_inside = make_paragraph("TextBox 2: Inside wps:wsp > wps:txbx (group shape)")
    w_txbx.append(p_inside)
    r_elem3.append(alt)

    # --- Section 4: Text box with multiple paragraphs inside ---
    p4 = doc.add_paragraph()
    run4 = p4.add_run()
    r_elem4 = run4._element
    txbx4 = make_textbox_paragraph("TextBox 3A: First paragraph in multi-para textbox")
    txbx4.append(make_paragraph("TextBox 3B: Second paragraph in same textbox"))
    r_elem4.append(txbx4)

    # --- Section 5: Text box in header ---
    header = doc.sections[0].header
    h_para = header.paragraphs[0]  # python-docx creates an empty paragraph
    h_run = h_para.add_run()
    h_txbx = make_textbox_paragraph("TextBox 4: Text box in header")
    h_run._element.append(h_txbx)

    # --- Section 6: Deeply nested text box (multiple group shape layers) ---
    p6 = doc.add_paragraph()
    run6 = p6.add_run()
    r_elem6 = run6._element

    alt2 = etree.Element(f'{{{MC}}}AlternateContent')
    choice2 = etree.SubElement(alt2, f'{{{MC}}}Choice', Requires='wps')
    wsp2 = etree.SubElement(choice2, f'{{{WPS}}}wsp')

    # Embed a second level: another wsp inside the first
    inner_wsp = etree.SubElement(wsp2, f'{{{WPS}}}wsp')
    inner_txbx = etree.SubElement(inner_wsp, f'{{{WPS}}}txbx')
    inner_w_txbx = etree.SubElement(inner_txbx, f'{W}txbxContent')
    inner_w_txbx.append(make_paragraph("TextBox 5: Deeply nested (2-level group shape)"))
    r_elem6.append(alt2)

    # --- Section 7: Legacy v:textbox pattern (simulated) ---
    V = 'urn:schemas-microsoft-com:vml'
    p7 = doc.add_paragraph()
    run7 = p7.add_run()
    r_elem7 = run7._element
    # Create a legacy-style v:textbox
    v_txbx = etree.Element(f'{{{V}}}textbox')
    legacy_txbx = make_textbox_paragraph("TextBox 6: Legacy v:textbox style")
    v_txbx.append(legacy_txbx)
    r_elem7.append(v_txbx)

    # Save
    output_dir = os.path.dirname(__file__)
    output_path = os.path.join(output_dir, 'sample_textboxes.docx')
    doc.save(output_path)
    print(f"Created {output_path}")
    print("Text boxes created:")
    print("  1. Direct text box in body")
    print("  2. Inside wps:wsp > wps:txbx (group shape)")
    print("  3A/3B. Multi-paragraph text box")
    print("  4. Text box in header")
    print("  5. Deeply nested (2-level group shape)")
    print("  6. Legacy v:textbox style")

if __name__ == '__main__':
    main()
