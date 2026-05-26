"""Generate a DOCX with an image placeholder to test image preservation."""
from docx import Document
from docx.shared import Inches
from lxml import etree
import os

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

doc = Document()
p = doc.add_paragraph()

# Run 1: text
run1 = p.add_run('Text before image ')

# Inline image placeholder (drawing element)
drawing = etree.SubElement(run1._element, f'{{{W}}}drawing')
inline = etree.SubElement(drawing, f'{{{WP}}}inline')
extent = etree.SubElement(inline, f'{{{WP}}}extent')
extent.set('cx', '914400'); extent.set('cy', '914400')

graphic = etree.SubElement(inline, f'{{{A}}}graphic')
gd = etree.SubElement(graphic, f'{{{A}}}graphicData')
gd.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
pic = etree.SubElement(gd, f'{{{A}}}pic')
nvPicPr = etree.SubElement(pic, f'{{{A}}}nvPicPr')
cNvPr = etree.SubElement(nvPicPr, f'{{{A}}}cNvPr')
cNvPr.set('id', '0'); cNvPr.set('name', 'Image Placeholder')
blipFill = etree.SubElement(pic, f'{{{A}}}blipFill')
blip = etree.SubElement(blipFill, f'{{{A}}}blip')
blip.set(f'{{{R}}}embed', 'rId1')
spPr = etree.SubElement(pic, f'{{{A}}}spPr')
prstGeom = etree.SubElement(spPr, f'{{{A}}}prstGeom')
prstGeom.set('prst', 'rect')

# Run 2: more text
p.add_run(' text after image')

# Another text box paragraph
doc.add_paragraph('Second paragraph with no image')

output_dir = os.path.dirname(__file__)
path = os.path.join(output_dir, 'sample_images.docx')
doc.save(path)
print(f'Created {path}')
