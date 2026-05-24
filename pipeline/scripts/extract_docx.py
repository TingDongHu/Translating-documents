import argparse
import json
from pathlib import Path
from docx import Document

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': W_NS}


def paragraph_text(paragraph):
    return ''.join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text


def add_entry(entries, tag, typ, text, **meta):
    entry = {'tag': tag, 'type': typ, 'text': text}
    entry.update(meta)
    entries.append(entry)


def extract_textboxes(doc, entries, next_index):
    xml_bodies = [doc.element.body]
    for section in doc.sections:
        for hf in [section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer]:
            if hf is not None:
                xml_bodies.append(hf._element)
    textbox_index = 0
    for xml_body in xml_bodies:
        for txbx in xml_body.findall('.//w:txbxContent', NSMAP):
            for para_index, paragraph in enumerate(txbx.findall(f'{{{W_NS}}}p')):
                text = ''.join(t.text or '' for t in paragraph.findall(f'.//{{{W_NS}}}t'))
                tag = f'P{next_index}'
                add_entry(entries, tag, 'textbox', text, txbx_index=textbox_index, para_index=para_index)
                next_index += 1
            textbox_index += 1
    return next_index


def main():
    parser = argparse.ArgumentParser(description='Extract tagged text from DOCX.')
    parser.add_argument('--source', required=True)
    parser.add_argument('--extraction-output', required=True)
    parser.add_argument('--tagged-output', required=True)
    parser.add_argument('--report-output', required=True)
    args = parser.parse_args()

    doc = Document(args.source)
    entries = []
    index = 0

    for para_index, paragraph in enumerate(doc.paragraphs):
        add_entry(entries, f'P{index}', 'body', paragraph_text(paragraph), para_index=para_index)
        index += 1

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_index, paragraph in enumerate(cell.paragraphs):
                    add_entry(entries, f'P{index}', 'body_table', paragraph_text(paragraph), table_idx=table_idx, row=row_idx, col=col_idx, para_index=para_index)
                    index += 1

    for section_idx, section in enumerate(doc.sections):
        for attr in ['header', 'first_page_header', 'even_page_header']:
            hf = getattr(section, attr)
            for para_index, paragraph in enumerate(hf.paragraphs):
                add_entry(entries, f'P{index}', 'header', paragraph_text(paragraph), section_idx=section_idx, header_type=attr, para_index=para_index)
                index += 1
            for table_idx, table in enumerate(hf.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for para_index, paragraph in enumerate(cell.paragraphs):
                            add_entry(entries, f'P{index}', 'header_table', paragraph_text(paragraph), section_idx=section_idx, header_type=attr, table_idx=table_idx, row=row_idx, col=col_idx, para_index=para_index)
                            index += 1
        for attr in ['footer', 'first_page_footer', 'even_page_footer']:
            hf = getattr(section, attr)
            for para_index, paragraph in enumerate(hf.paragraphs):
                add_entry(entries, f'P{index}', 'footer', paragraph_text(paragraph), section_idx=section_idx, footer_type=attr, para_index=para_index)
                index += 1
            for table_idx, table in enumerate(hf.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for para_index, paragraph in enumerate(cell.paragraphs):
                            add_entry(entries, f'P{index}', 'footer_table', paragraph_text(paragraph), section_idx=section_idx, footer_type=attr, table_idx=table_idx, row=row_idx, col=col_idx, para_index=para_index)
                            index += 1

    index = extract_textboxes(doc, entries, index)

    extraction = {'source': args.source, 'paragraphs': entries}
    Path(args.extraction_output).parent.mkdir(parents=True, exist_ok=True)
    Path(args.extraction_output).write_text(json.dumps(extraction, ensure_ascii=False, indent=2), encoding='utf-8')
    Path(args.tagged_output).write_text('\n'.join(f"[{entry['tag']}] {entry['text']}" for entry in entries) + '\n', encoding='utf-8')
    report = {
        'paragraph_count': len(entries),
        'body_count': sum(1 for entry in entries if entry['type'] == 'body'),
        'table_count': sum(1 for entry in entries if entry['type'] == 'body_table'),
        'header_count': sum(1 for entry in entries if entry['type'].startswith('header')),
        'footer_count': sum(1 for entry in entries if entry['type'].startswith('footer')),
        'textbox_count': sum(1 for entry in entries if entry['type'] == 'textbox'),
        'warnings': [],
    }
    Path(args.report_output).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"Extracted {len(entries)} tagged entries")


if __name__ == '__main__':
    main()
